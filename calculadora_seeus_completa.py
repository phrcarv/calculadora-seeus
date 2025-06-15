
import streamlit as st
from datetime import datetime, timedelta
from numpy import busday_count
import pandas as pd
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Calculadora", layout="centered")
st.title("📊 Calculadora")

feriados = pd.to_datetime([
    "1973-01-01", "1973-04-21", "1973-05-01", "1973-09-07", "1973-10-12", "1973-11-02", "1973-11-15", "1973-12-25",
    "2024-01-01", "2024-04-21", "2024-05-01", "2024-09-07", "2024-10-12", "2024-11-02", "2024-11-15", "2024-12-25",
    "2025-01-01", "2025-04-21", "2025-05-01", "2025-09-07", "2025-10-12", "2025-11-02", "2025-11-15", "2025-12-25",
    "2026-01-01", "2026-04-21", "2026-05-01", "2026-09-07", "2026-10-12", "2026-11-02", "2026-11-15", "2026-12-25"
]).date

tab1, tab2, tab3 = st.tabs(["📅 Recolhimento Noturno", "📚 Remição por Estudo (em breve)", "💰 Dias-Multa"])

with tab1:
    
    st.markdown("### Cálculo de Detração por Recolhimento Noturno")

    with st.form("form_periodos"):
        n = st.number_input("Quantos períodos deseja calcular?", min_value=1, max_value=10, value=1)
        horas_noite = st.number_input("Horas por noite (em dias úteis)", min_value=0.0, max_value=24.0, step=0.5, value=8.0)
        modo = st.selectbox("Tratamento dos fins de semana", ["Tratar como dias úteis", "Recolhimento integral (24h)", "Ignorar (somente dias úteis)"])
        dias_corridos = st.radio("Cálculo por dias corridos?", ["Sim", "Não"])
        tratamento_feriados = st.selectbox("Tratamento dos feriados", ["Ignorar", "Considerar como dias úteis", "Recolhimento integral (24h)"])

        datas = []
        for i in range(n):
            st.markdown(f"#### Período {i+1}")
            col1, col2 = st.columns(2)
            with col1:
                inicio = st.date_input(f"Início {i+1}", key=f"inicio_{i}")
            with col2:
                fim = st.date_input(f"Fim {i+1}", key=f"fim_{i}")
            datas.append((inicio, fim))

        submitted = st.form_submit_button("Calcular Detração")

    if submitted:
        total_dias = 0
        total_horas = 0
        total_dias_uteis = 0
        total_dias_fds = 0
        total_feriados = 0
        total_horas_uteis = 0
        total_horas_fds = 0
        total_horas_feriados = 0
        relatorio_detalhado = []

        for i, (inicio, fim) in enumerate(datas):
            if inicio > fim:
                st.error(f"O início do Período {i+1} não pode ser após o fim.")
                continue
            dias = (fim - inicio).days + 1
            if dias_corridos == "Sim":
                horas = dias * horas_noite
                relatorio_detalhado.append((i+1, inicio, fim, dias, horas, "dias corridos"))
            else:
                feriados_intervalo = [f for f in feriados if inicio <= f <= fim]
                if tratamento_feriados == "Ignorar":
                    feriados_utilizados = []
                else:
                    feriados_utilizados = feriados_intervalo

                dias_uteis = busday_count(inicio.isoformat(), (fim + timedelta(days=1)).isoformat(), holidays=feriados_utilizados)
                dias_fds = dias - dias_uteis

                horas_uteis = dias_uteis * horas_noite
                horas_fds = dias_fds * (24 if modo == "Recolhimento integral (24h)" else (horas_noite if modo == "Tratar como dias úteis" else 0))
                horas_feriados = len(feriados_intervalo) * (24 if tratamento_feriados == "Recolhimento integral (24h)" else (horas_noite if tratamento_feriados == "Considerar como dias úteis" else 0))

                horas = horas_uteis + horas_fds + horas_feriados

                total_dias_uteis += dias_uteis
                total_dias_fds += dias_fds
                total_feriados += len(feriados_intervalo)
                total_horas_uteis += horas_uteis
                total_horas_fds += horas_fds
                total_horas_feriados += horas_feriados

                relatorio_detalhado.append((i+1, inicio, fim, dias, horas, f"{dias_uteis} úteis, {dias_fds} fds, {len(feriados_intervalo)} feriados"))
            total_dias += dias
            total_horas += horas

        dias_detracao = total_horas / 24

        st.success(f"Detração total estimada: {dias_detracao:.2f} dias")

        st.markdown("### 🧾 Detalhamento:")
        st.markdown(f"- Dias úteis: {total_dias_uteis} dias — {total_horas_uteis:.2f} horas")
        st.markdown(f"- Finais de semana: {total_dias_fds} dias — {total_horas_fds:.2f} horas")
        st.markdown(f"- Feriados: {total_feriados} dias — {total_horas_feriados:.2f} horas")

        texto_seeus = f"""Em cumprimento à decisão exarada nos autos, procedo às seguintes anotações:\nO recuperando permaneceu em recolhimento noturno do dia {datas[0][0].strftime('%d/%m/%Y')} a {datas[-1][1].strftime('%d/%m/%Y')}.\nNo período indicado, cumpriu {total_horas:.2f} horas de recolhimento, correspondentes a {dias_detracao:.2f} dias de detração."""
        st.text_area("📋 Texto para certificar no SEEUs:", value=texto_seeus, height=150)

        doc = Document()
        doc.add_heading('Relatório de Detração', 0)
        doc.add_paragraph(f"Tratamento dos fins de semana: {modo}")
        doc.add_paragraph(f"Tratamento dos feriados: {tratamento_feriados}")
        doc.add_paragraph(f"Cálculo por dias corridos: {dias_corridos}")
        doc.add_paragraph(f"Horas por noite (dias úteis): {horas_noite}h")

        for item in relatorio_detalhado:
            doc.add_paragraph(f"Período {item[0]}: {item[1].strftime('%d/%m/%Y')} a {item[2].strftime('%d/%m/%Y')} ({item[3]} dias) - {item[4]:.2f} horas computadas ({item[5]})")

        doc.add_paragraph(f"\nTotal de horas: {total_horas:.2f}")
        doc.add_paragraph(f"Equivalente em dias de detração: {dias_detracao:.2f} dias")
        doc.add_paragraph(f"Dias úteis: {total_dias_uteis} dias — {total_horas_uteis:.2f} horas")
        doc.add_paragraph(f"Finais de semana: {total_dias_fds} dias — {total_horas_fds:.2f} horas")
        doc.add_paragraph(f"Feriados: {total_feriados} dias — {total_horas_feriados:.2f} horas")
        doc.add_paragraph("Criado por Pedro Henrique Ribeiro de Carvalho")
        doc.add_paragraph("Gestor da 1ª Vara Criminal de Água Boa - MT")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("📄 Baixar relatório em Word", buffer, file_name="relatorio_detracao.docx")


with tab2:
    st.markdown("### Em breve:")
    st.info("A calculadora de **remição por estudo** será incluída nesta aba com base na norma vigente. Aguarde as próximas atualizações.")

with tab3:
    st.markdown("### 💰 Cálculo de Dias-Multa")

    salario_atual = 1412.00
    usar_salario_atual = st.checkbox("Usar salário mínimo atual (R$ 1.412,00)?", value=True)

    if usar_salario_atual:
        salario = salario_atual
        data_base = datetime(2024, 1, 1)
    else:
        salario = st.number_input("Salário mínimo da época (R$):", min_value=0.01)
        data_base = st.date_input("Data-base do salário informado:")

    dias_multa = st.number_input("Quantidade de dias-multa fixada:", min_value=1)
    fracao = st.selectbox("Fração do salário mínimo por dia-multa:", ["1/30", "1/50", "1/10", "Outro"])

    if fracao == "Outro":
        denominador = st.number_input("Informe o denominador da fração (ex: para 1/45, digite 45):", min_value=1)
    else:
        denominador = int(fracao.split("/")[1])

    data_referencia = st.date_input("Data final para correção monetária:", value=datetime.today())
    anos = max(0, data_referencia.year - data_base.year)
    ipca_medio = 0.05
    fator_correcao = (1 + ipca_medio) ** anos

    valor_dia_original = salario / denominador
    valor_dia_corrigido = valor_dia_original * fator_correcao
    valor_total_corrigido = valor_dia_corrigido * dias_multa

    st.markdown(f"**Valor original do dia-multa:** R$ {valor_dia_original:.2f}")
    st.markdown(f"**Valor corrigido do dia-multa:** R$ {valor_dia_corrigido:.2f}")
    st.markdown(f"**Total corrigido a pagar:** R$ {valor_total_corrigido:.2f}")

    texto_certidao = f"""Nos termos da sentença condenatória, considerando o salário mínimo vigente à época (R$ {salario:.2f}) e a fração de 1/{denominador}, o valor do dia-multa foi fixado em R$ {valor_dia_original:.2f}.
Atualizado monetariamente pelo índice IPCA-E até {data_referencia.strftime('%d/%m/%Y')}, o valor do dia-multa é de R$ {valor_dia_corrigido:.2f}, totalizando R$ {valor_total_corrigido:.2f} pelos {int(dias_multa)} dias-multa."""

    st.text_area("📋 Texto para certidão:", value=texto_certidao, height=180)

st.markdown("""
<hr style="margin-top: 3em; margin-bottom: 1em">
<p style='text-align: center; font-size: 0.85em; color: gray;'>
Criado por <strong>Pedro Henrique Ribeiro de Carvalho</strong><br>
Gestor da 1ª Vara Criminal de Água Boa - MT
</p>
""", unsafe_allow_html=True)
